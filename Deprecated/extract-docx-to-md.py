"""Extract all text from a docx to a markdown file."""
import sys
from pathlib import Path
from docx import Document

path = Path(sys.argv[1])
doc = Document(str(path))
out_path = path.with_suffix(".extracted.md")

lines = []
lines.append(f"# Extracted docx: {path.name}")
lines.append(f"")
lines.append(f"Paragraph count: {len(doc.paragraphs)}")
lines.append(f"Table count: {len(doc.tables)}")
lines.append(f"")
lines.append(f"---")
lines.append(f"")

for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        lines.append("")
        continue
    style = para.style.name if para.style else ""
    if style.startswith("Heading"):
        level = style.replace("Heading ", "").strip() or "1"
        try:
            n = int(level)
            prefix = "#" * min(n + 1, 6)
            lines.append(f"{prefix} {text}")
        except Exception:
            lines.append(f"## {text}")
    else:
        lines.append(text)
    lines.append("")

for i, table in enumerate(doc.tables, 1):
    lines.append(f"")
    lines.append(f"### Table {i}")
    lines.append(f"")
    for row in table.rows:
        row_text = " | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells)
        lines.append(f"| {row_text} |")
    lines.append("")

out_path.write_text("\n".join(lines), encoding="utf-8")
print(f"Wrote {out_path} ({out_path.stat().st_size} bytes, {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables)")
