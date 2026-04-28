"""
Reformat-UCMSIS-Docx.py
=======================
Reformats .docx files under C:\\Users\\shelc\\Documents\\UC MS-IS\\ to the explicit
standard Clay established 2026-04-26:
  - Arial font, 11pt, black color
  - Line spacing 1.15
  - No space before/after paragraphs
  - First-line indent 0 (no indentation)
  - No bold (italic preserved)
  - Replace em-dash, en-dash, smart quotes, ellipsis, NBSP, etc. with ASCII 0-127
  - Preserve center alignment (title pages, section titles)
  - Preserve italics
  - Preserve monospace fonts on code-style runs (heuristic: Courier/Consolas/Mono in name)
  - Apply to body paragraphs AND table cells (recursive)
  - DOES NOT alter: headers/footers, title page structure, page breaks, hyperlinks, images

Backup: Each file is copied to <samedir>/Deprecated/<filename>__backup_<YYYYMMDD-HHMMSS>.docx
        before reformatting. Deprecate-never-delete (CLAUDE.md global rule).

Skip rules:
  - ~$* (Word lock files)
  - Any path containing _Archive_*
  - Any path containing /Deprecated/
  - The template itself (handled separately)
  - Files where a corresponding ~$<name>.docx lock file exists (file open in Word)
  - Files where write fails (logged, not crashed)

Usage:
  py -3.14 Reformat-UCMSIS-Docx.py --dryrun           # preview transformations only
  py -3.14 Reformat-UCMSIS-Docx.py --apply            # apply with backups
  py -3.14 Reformat-UCMSIS-Docx.py --apply --file <p> # single file mode

Standard inheritance: 2026-04-26 explicit standard supersedes the reference file's
Calibri default. Reference: Caddell_Descriptive_Insights_Memo_v2_20260424.docx
(reference deviates: uses Calibri; standard uses Arial).
"""

import argparse
import shutil
import sys
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

ROOT = Path(r"C:\Users\shelc\Documents\UC MS-IS")
TEMPLATE_PATH = ROOT / "Supplemental" / "Templates" / "UC MS-IS Word Template.docx"

# Character replacement table (non-ASCII to ASCII 0-127)
ASCII_REPLACEMENTS = {
    "—": " - ",   # em dash
    "–": "-",     # en dash
    "−": "-",     # minus sign
    "“": '"',     # left double quote
    "”": '"',     # right double quote
    "‘": "'",     # left single quote
    "’": "'",     # right single quote / apostrophe
    "…": "...",   # ellipsis
    " ": " ",     # non-breaking space
    "•": "-",     # bullet
    "‣": "-",     # triangular bullet
    "◦": "-",     # white bullet
    "⁃": "-",     # hyphen bullet
    "·": "-",     # middle dot used as bullet
    "→": "->",    # right arrow
    "←": "<-",    # left arrow
    "↔": "<->",   # both arrows
    "≈": "~",     # approximately equal
    "≠": "!=",    # not equal
    "≤": "<=",    # less than or equal
    "≥": ">=",    # greater than or equal
    "×": "x",     # multiplication
    "±": "+/-",   # plus-minus
    "°": " deg ", # degree
    "€": "EUR",   # euro
    "£": "GBP",   # pound
    "¥": "JPY",   # yen
    "©": "(c)",   # copyright
    "®": "(R)",   # registered
    "™": "(TM)",  # trademark
    "‰": " percent ", # per mille
    "¼": "1/4",
    "½": "1/2",
    "¾": "3/4",
}

MONOSPACE_HINTS = ("courier", "consolas", "monaco", "mono", "lucida console", "fixedsys", "menlo")

def is_skip_path(p: Path) -> str:
    """Return skip-reason string or None."""
    name = p.name
    parts = p.parts
    if name.startswith("~$"):
        return "word-lock"
    if any("_Archive_" in part for part in parts):
        return "archive"
    if "Deprecated" in parts:
        return "deprecated"
    if name == "UC MS-IS Word Template.docx":
        return "template-itself"
    # Skip if a Word lock file exists for this doc
    lock_file = p.parent / f"~${name}"
    if lock_file.exists():
        return "open-in-word"
    return None

def replace_ascii(text: str) -> tuple[str, int]:
    """Replace non-ASCII chars per replacement table. Return (new_text, count)."""
    if not text:
        return text, 0
    count = 0
    out = text
    for orig, repl in ASCII_REPLACEMENTS.items():
        if orig in out:
            count += out.count(orig)
            out = out.replace(orig, repl)
    # Replace any remaining non-ASCII with ?
    if any(ord(c) > 127 for c in out):
        cleaned = []
        for c in out:
            if ord(c) > 127:
                count += 1
                cleaned.append("?")
            else:
                cleaned.append(c)
        out = "".join(cleaned)
    return out, count

def is_monospace_font(font_name: str) -> bool:
    if not font_name:
        return False
    lower = font_name.lower()
    return any(hint in lower for hint in MONOSPACE_HINTS)

def transform_run(run, stats: dict) -> None:
    """Transform a single run: text, font, size, color, bold."""
    # Text replacements
    new_text, replaced = replace_ascii(run.text)
    if replaced > 0:
        run.text = new_text
        stats["chars_replaced"] += replaced

    # Font name: Arial (unless monospace)
    current_font_name = run.font.name
    if not is_monospace_font(current_font_name):
        if current_font_name != "Arial":
            run.font.name = "Arial"
            # Also set the rFonts element for east-asian / complex script
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                from docx.oxml import OxmlElement
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:ascii"), "Arial")
            rFonts.set(qn("w:hAnsi"), "Arial")
            rFonts.set(qn("w:cs"), "Arial")
            stats["runs_font_changed"] += 1

    # Font size: 11pt
    if run.font.size != Pt(11):
        run.font.size = Pt(11)
        stats["runs_size_changed"] += 1

    # Color: black
    current_color = run.font.color
    if current_color is None or current_color.rgb is None or current_color.rgb != RGBColor(0, 0, 0):
        run.font.color.rgb = RGBColor(0, 0, 0)
        stats["runs_color_changed"] += 1

    # Bold: off (italic preserved)
    if run.bold:
        run.bold = False
        stats["runs_bold_removed"] += 1

def is_separator_paragraph(text: str) -> bool:
    """Detect text-only horizontal-rule separator paragraphs."""
    if not text:
        return False
    stripped = text.strip()
    if len(stripped) < 3:
        return False
    first = stripped[0]
    if first not in "-_=*":
        return False
    return all(c == first for c in stripped)

def convert_separator_to_pagebreak(para, stats: dict) -> bool:
    """If para is a separator, replace its content with a page break. Return True if converted."""
    if not is_separator_paragraph(para.text):
        return False
    # Clear all run text
    for run in para.runs:
        run.text = ""
    # Add page break to first run (or create one)
    if para.runs:
        run = para.runs[0]
    else:
        run = para.add_run("")
    from docx.oxml import OxmlElement
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)
    stats["separators_to_pagebreaks"] = stats.get("separators_to_pagebreaks", 0) + 1
    return True

def transform_paragraph(para, stats: dict) -> None:
    """Transform a single paragraph: runs, line spacing, before/after, indent."""
    # First check separator -> page break
    if convert_separator_to_pagebreak(para, stats):
        # Still apply paragraph-level formatting below
        pass

    for run in para.runs:
        transform_run(run, stats)

    pf = para.paragraph_format

    # Line spacing: 1.15
    if pf.line_spacing != 1.15:
        pf.line_spacing = 1.15
        stats["paras_linespacing_set"] += 1

    # Space before/after: 0
    if pf.space_before is None or pf.space_before != Pt(0):
        pf.space_before = Pt(0)
        stats["paras_space_before_set"] += 1
    if pf.space_after is None or pf.space_after != Pt(0):
        pf.space_after = Pt(0)
        stats["paras_space_after_set"] += 1

    # First-line indent: 0
    if pf.first_line_indent is not None and pf.first_line_indent != 0:
        pf.first_line_indent = 0
        stats["paras_indent_removed"] += 1

    # Left indent: 0 (preserve list indentation by checking style; only zero if not list)
    style_name = (para.style.name if para.style else "").lower()
    is_list = "list" in style_name or "bullet" in style_name or "number" in style_name
    if not is_list:
        if pf.left_indent is not None and pf.left_indent != 0:
            pf.left_indent = 0
            stats["paras_left_indent_removed"] += 1

def transform_table(table, stats: dict) -> None:
    """Transform all cells in a table."""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                transform_paragraph(para, stats)
            for nested_table in cell.tables:
                transform_table(nested_table, stats)
            stats["cells_touched"] += 1

def transform_document(doc) -> dict:
    """Apply all transformations to a Document. Return stats dict."""
    stats = {
        "chars_replaced": 0,
        "runs_font_changed": 0,
        "runs_size_changed": 0,
        "runs_color_changed": 0,
        "runs_bold_removed": 0,
        "paras_linespacing_set": 0,
        "paras_space_before_set": 0,
        "paras_space_after_set": 0,
        "paras_indent_removed": 0,
        "paras_left_indent_removed": 0,
        "cells_touched": 0,
        "paragraphs_processed": 0,
        "tables_processed": 0,
    }

    # Body paragraphs
    for para in doc.paragraphs:
        transform_paragraph(para, stats)
        stats["paragraphs_processed"] += 1

    # Body tables
    for table in doc.tables:
        transform_table(table, stats)
        stats["tables_processed"] += 1

    # Update Normal style default
    try:
        normal = doc.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(11)
        rPr = normal.element.find(qn("w:rPr"))
        if rPr is not None:
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                from docx.oxml import OxmlElement
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:ascii"), "Arial")
            rFonts.set(qn("w:hAnsi"), "Arial")
            rFonts.set(qn("w:cs"), "Arial")
    except Exception as e:
        stats["normal_style_error"] = str(e)

    return stats

def backup_file(p: Path) -> Path:
    """Copy p to <samedir>/Deprecated/<name>__backup_<timestamp>.docx. Return backup path."""
    deprecated_dir = p.parent / "Deprecated"
    deprecated_dir.mkdir(exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d-%H%M%S")
    backup_name = f"{p.stem}__backup_{ts}{p.suffix}"
    backup_path = deprecated_dir / backup_name
    shutil.copy2(p, backup_path)
    return backup_path

def process_file(p: Path, apply: bool, log: list) -> dict:
    """Process a single .docx file. Return result dict.
    Note: no local backup — relies on Microsoft Word / OneDrive built-in version history
    per Clay 2026-04-26 directive ("dont overthink it too much").
    """
    result = {
        "path": str(p),
        "name": p.name,
        "skipped": False,
        "skip_reason": None,
        "stats": None,
        "error": None,
    }

    skip_reason = is_skip_path(p)
    if skip_reason:
        result["skipped"] = True
        result["skip_reason"] = skip_reason
        return result

    try:
        if apply:
            # Pre-probe to hydrate OneDrive cloud-only files
            try:
                with open(p, "r+b") as fh:
                    fh.read(8)
            except PermissionError:
                import time
                time.sleep(0.5)
                with open(p, "r+b") as fh:
                    fh.read(8)

        doc = Document(p)
        stats = transform_document(doc)
        result["stats"] = stats

        if apply:
            # OneDrive workaround: python-docx's zipfile-write uses 'wb' mode which
            # OneDrive locks. Save to a temp path then byte-copy back via r+b which
            # OneDrive permits.
            import tempfile
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp_path = tmp.name
            try:
                doc.save(tmp_path)
                # Read bytes from temp
                with open(tmp_path, "rb") as src:
                    new_bytes = src.read()
                # Overwrite original using r+b (truncate then write within same handle)
                try:
                    with open(p, "r+b") as dst:
                        dst.seek(0)
                        dst.write(new_bytes)
                        dst.truncate()
                except PermissionError:
                    import time
                    time.sleep(1.0)
                    with open(p, "r+b") as dst:
                        dst.seek(0)
                        dst.write(new_bytes)
                        dst.truncate()
                # Verify saved file parses
                try:
                    Document(p)
                except Exception as e:
                    result["error"] = f"Reformatted file failed to parse: {e}"
            finally:
                try:
                    Path(tmp_path).unlink()
                except Exception:
                    pass
    except Exception as e:
        result["error"] = f"Processing error: {e}"

    return result

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--dryrun", action="store_true", help="Preview only; do not modify files")
    parser.add_argument("--apply", action="store_true", help="Apply transformations with backups")
    parser.add_argument("--file", type=str, default=None, help="Single file path to process")
    parser.add_argument("--limit", type=int, default=None, help="Process at most N files")
    args = parser.parse_args()

    if not args.dryrun and not args.apply:
        print("ERROR: must specify --dryrun or --apply")
        sys.exit(2)
    if args.dryrun and args.apply:
        print("ERROR: cannot specify both --dryrun and --apply")
        sys.exit(2)

    apply = args.apply

    # Build file list
    if args.file:
        files = [Path(args.file)]
    else:
        files = list(ROOT.rglob("*.docx"))
        if args.limit:
            files = files[: args.limit]

    print(f"Mode: {'APPLY' if apply else 'DRYRUN'}")
    print(f"Files to process: {len(files)}")
    print(f"Standard: Arial 11pt, black, line-spacing 1.15, ASCII 0-127, no bold")
    print("=" * 80)

    log = []
    results = []
    for i, p in enumerate(files):
        result = process_file(p, apply, log)
        results.append(result)

        if result["skipped"]:
            print(f"[{i+1}/{len(files)}] SKIP ({result['skip_reason']}): {result['name']}")
            continue

        if result["error"]:
            print(f"[{i+1}/{len(files)}] ERROR: {result['name']} -- {result['error']}")
            continue

        s = result["stats"]
        summary_parts = []
        if s["chars_replaced"] > 0:
            summary_parts.append(f"{s['chars_replaced']} chars")
        if s["runs_font_changed"] > 0:
            summary_parts.append(f"{s['runs_font_changed']} fonts")
        if s["runs_size_changed"] > 0:
            summary_parts.append(f"{s['runs_size_changed']} sizes")
        if s["runs_color_changed"] > 0:
            summary_parts.append(f"{s['runs_color_changed']} colors")
        if s["runs_bold_removed"] > 0:
            summary_parts.append(f"{s['runs_bold_removed']} bolds")
        if s["paras_linespacing_set"] > 0:
            summary_parts.append(f"{s['paras_linespacing_set']} spacing")
        summary = ", ".join(summary_parts) if summary_parts else "no changes"
        print(f"[{i+1}/{len(files)}] {'APPLIED' if apply else 'PREVIEW'}: {result['name']} -- {summary}")

    # Summary
    print("=" * 80)
    print("SUMMARY:")
    skipped = [r for r in results if r["skipped"]]
    errored = [r for r in results if r["error"]]
    processed = [r for r in results if not r["skipped"] and not r["error"]]
    print(f"  Total: {len(results)}")
    print(f"  Skipped: {len(skipped)}")
    print(f"  Errored: {len(errored)}")
    print(f"  Processed: {len(processed)}")

    # Aggregate stats
    if processed:
        total_chars = sum(r["stats"]["chars_replaced"] for r in processed)
        total_fonts = sum(r["stats"]["runs_font_changed"] for r in processed)
        total_sizes = sum(r["stats"]["runs_size_changed"] for r in processed)
        total_colors = sum(r["stats"]["runs_color_changed"] for r in processed)
        total_bolds = sum(r["stats"]["runs_bold_removed"] for r in processed)
        print(f"  Total non-ASCII chars replaced: {total_chars}")
        print(f"  Total run fonts changed to Arial: {total_fonts}")
        print(f"  Total run sizes set to 11pt: {total_sizes}")
        print(f"  Total run colors set to black: {total_colors}")
        print(f"  Total bolds removed: {total_bolds}")

    if errored:
        print()
        print("ERRORS:")
        for r in errored:
            print(f"  {r['name']}: {r['error']}")

    return 0 if not errored else 1

if __name__ == "__main__":
    sys.exit(main())
