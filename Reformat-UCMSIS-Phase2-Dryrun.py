"""
Reformat-UCMSIS-Phase2-Dryrun.py
================================
Phase 2 dry-run analysis for the UC MS-IS reformatter. Reports per-file:

1. SHAPES inventory: pictures (preserve), charts (preserve), other shapes (would propose remove)
2. TITLE PAGE state: detect what's currently in the first ~25 lines and compare to spec
3. HEADER state: current header content vs spec (class code & instance left, assignment name right)
4. FOOTER state: current footer content vs spec (UC Lindner insignia + page number)

Modifies nothing. Outputs JSON + summary table to stdout.
"""

import json
import sys
from pathlib import Path
from collections import Counter
from docx import Document
from docx.oxml.ns import qn

ROOT = Path(r"C:\Users\shelc\Documents\UC MS-IS")

# OOXML namespace URIs
NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
    "c": "http://schemas.openxmlformats.org/drawingml/2006/chart",
    "dgm": "http://schemas.openxmlformats.org/drawingml/2006/diagram",
}

ARTIFACT_URIS = {
    "http://schemas.openxmlformats.org/drawingml/2006/picture": "picture",
    "http://schemas.openxmlformats.org/drawingml/2006/chart": "chart",
    "http://schemas.openxmlformats.org/drawingml/2006/diagram": "diagram",
}

DECORATIVE_URIS = {
    "http://schemas.microsoft.com/office/word/2010/wordprocessingShape": "shape",
    "http://schemas.microsoft.com/office/drawing/2014/chartex": "chart-ex",
}

def is_skip(p):
    if p.name.startswith("~$"): return "word-lock"
    if any("_Archive_" in part for part in p.parts): return "archive"
    if "Deprecated" in p.parts: return "deprecated"
    if p.name == "UC MS-IS Word Template.docx": return "template-itself"
    if (p.parent / f"~${p.name}").exists(): return "open-in-word"
    return None

def analyze_shapes(doc):
    """Walk the body XML and inventory all <w:drawing> elements; categorize by graphicData URI."""
    body = doc.element.body
    drawings = body.findall(f".//{{{NS['w']}}}drawing")
    by_kind = Counter()
    drawing_details = []
    for d in drawings:
        gd = d.find(f".//{{{NS['a']}}}graphicData")
        uri = gd.get("uri") if gd is not None else None
        if uri in ARTIFACT_URIS:
            kind = ARTIFACT_URIS[uri]
        elif uri in DECORATIVE_URIS:
            kind = DECORATIVE_URIS[uri]
        else:
            kind = f"unknown:{uri}" if uri else "unknown"
        by_kind[kind] += 1
        drawing_details.append({"uri": uri, "kind": kind})
    return {
        "total_drawings": len(drawings),
        "by_kind": dict(by_kind),
        "decorative_count": sum(by_kind[k] for k in by_kind if k not in ARTIFACT_URIS.values()),
        "artifact_count": sum(by_kind[k] for k in by_kind if k in ARTIFACT_URIS.values()),
    }

def analyze_title_page(doc):
    """Look at the first ~25 paragraphs; detect title-page-like structure."""
    paras = doc.paragraphs[:25]
    text_paras = [(i, p) for i, p in enumerate(paras) if p.text.strip()]
    centered_count = sum(1 for _, p in text_paras if str(p.alignment) == "CENTER (1)")

    # Heuristic: title page has multiple centered paragraphs with assignment / name / institution / course info
    likely_title_page = centered_count >= 3

    # Extract first non-empty centered text lines (the candidate title-page content)
    title_page_text = []
    for i, p in text_paras[:10]:
        if str(p.alignment) == "CENTER (1)":
            title_page_text.append(p.text.strip())

    return {
        "likely_has_title_page": likely_title_page,
        "centered_text_paras_in_first_25": centered_count,
        "title_page_text_sample": title_page_text[:6],
    }

def analyze_header_footer(doc):
    """Inspect first section header and footer text content."""
    if not doc.sections:
        return {"sections": 0}
    s = doc.sections[0]
    header_text = ""
    footer_text = ""
    if s.header and s.header.paragraphs:
        header_text = " | ".join(p.text for p in s.header.paragraphs if p.text.strip())
    if s.footer and s.footer.paragraphs:
        footer_text = " | ".join(p.text for p in s.footer.paragraphs if p.text.strip())

    # Check if header looks like the spec (contains both a class code pattern and assignment hint)
    has_class_code_pattern = any(c.isdigit() for c in header_text) and (
        "IS" in header_text.upper() or "IT" in header_text.upper() or "BA" in header_text.upper() or "ACCT" in header_text.upper()
    )
    has_tab_separator = "\t" in (s.header.paragraphs[0].text if s.header.paragraphs else "")

    # Check if footer has page-number-like content (XML field)
    footer_has_page_field = False
    if s.footer:
        for p in s.footer.paragraphs:
            for run in p.runs:
                # Check XML for fldSimple or fldChar with PAGE field
                xml = run._element.xml
                if "PAGE" in xml or "fldSimple" in xml or "fldChar" in xml:
                    footer_has_page_field = True
                    break
            if footer_has_page_field: break

    return {
        "header_text": header_text[:200],
        "footer_text": footer_text[:200],
        "header_has_class_code_pattern": has_class_code_pattern,
        "header_has_tab_separator": has_tab_separator,
        "footer_has_page_field": footer_has_page_field,
        "footer_empty": not footer_text and not footer_has_page_field,
    }

def analyze_file(p):
    """Run all phase-2 analyses on a single file."""
    result = {"path": str(p), "name": p.name}
    skip = is_skip(p)
    if skip:
        result["skipped"] = skip
        return result
    try:
        doc = Document(p)
        result["shapes"] = analyze_shapes(doc)
        result["title_page"] = analyze_title_page(doc)
        result["header_footer"] = analyze_header_footer(doc)
    except Exception as e:
        result["error"] = str(e)
    return result

def main():
    files = list(ROOT.rglob("*.docx"))
    results = []
    for p in files:
        results.append(analyze_file(p))

    # Summary
    skipped = [r for r in results if r.get("skipped")]
    errored = [r for r in results if r.get("error")]
    analyzed = [r for r in results if "shapes" in r]

    print(f"=== Phase 2 Dryrun Summary ===")
    print(f"Total files: {len(results)}")
    print(f"Skipped: {len(skipped)}")
    print(f"Errored: {len(errored)}")
    print(f"Analyzed: {len(analyzed)}")
    print()

    # Aggregate findings
    files_with_decorative_shapes = [r for r in analyzed if r["shapes"]["decorative_count"] > 0]
    files_without_title_page = [r for r in analyzed if not r["title_page"]["likely_has_title_page"]]
    files_with_title_page = [r for r in analyzed if r["title_page"]["likely_has_title_page"]]
    files_with_empty_footer = [r for r in analyzed if r["header_footer"].get("footer_empty")]
    files_with_header = [r for r in analyzed if r["header_footer"].get("header_text")]

    print(f"=== Shape Findings ===")
    print(f"Files with DECORATIVE shapes (would propose remove): {len(files_with_decorative_shapes)}")
    for r in files_with_decorative_shapes[:15]:
        kinds = r["shapes"]["by_kind"]
        decor = {k: v for k, v in kinds.items() if k not in ARTIFACT_URIS.values()}
        print(f"  {r['name']}: {decor}")
    if len(files_with_decorative_shapes) > 15:
        print(f"  ... and {len(files_with_decorative_shapes)-15} more")
    print()

    total_artifacts = sum(r["shapes"]["artifact_count"] for r in analyzed)
    total_decorative = sum(r["shapes"]["decorative_count"] for r in analyzed)
    print(f"Total artifacts (preserve): {total_artifacts}")
    print(f"Total decorative shapes (would remove): {total_decorative}")
    print()

    print(f"=== Title Page Findings ===")
    print(f"Files WITH likely title page: {len(files_with_title_page)}")
    print(f"Files WITHOUT title page (would propose adding): {len(files_without_title_page)}")
    print(f"  Examples (first 10):")
    for r in files_without_title_page[:10]:
        print(f"    {r['name']}")
    if len(files_without_title_page) > 10:
        print(f"  ... and {len(files_without_title_page)-10} more")
    print()

    print(f"=== Header/Footer Findings ===")
    print(f"Files with empty footer (would propose adding UC Lindner + page): {len(files_with_empty_footer)}")
    print(f"Files with existing header: {len(files_with_header)}")
    print(f"  Sample existing headers (first 10):")
    for r in files_with_header[:10]:
        print(f"    {r['name']}: {r['header_footer']['header_text'][:80]}")
    print()

    # Write full results JSON
    out_path = Path(r"C:\Users\shelc\Documents\Journal\Projects\Learning\UCMSIS_Phase2_Dryrun_20260426.json")
    out_path.parent.mkdir(exist_ok=True, parents=True)
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(results, f, indent=2)
    print(f"Full per-file results: {out_path}")

if __name__ == "__main__":
    main()
